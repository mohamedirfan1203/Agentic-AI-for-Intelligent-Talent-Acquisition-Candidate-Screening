"""
Final Report Agent — Compact 1-page Word report + email to HR
"""

import logging
import os
import smtplib
import threading
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from sqlalchemy.orm import Session

from app.models.tables import Candidate, InterviewChat, InterviewSession

load_dotenv()
logger = logging.getLogger("agents.final_report_agent")

# ── Config ────────────────────────────────────────────────────────────────────
REPORTS_DIR = Path(__file__).parent.parent.parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

HR_EMAIL    = os.getenv("HR_EMAIL", "")
SMTP_HOST   = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT   = int(os.getenv("SMTP_PORT", 587))

# Primary SMTP sender (report-specific, may be absent or use regular pw → falls back)
REPORT_FROM = os.getenv("REPORT_EMAIL_FROM", "")
REPORT_PASS = os.getenv("REPORT_EMAIL_PASS", "")

# Fallback: use the existing working app-password SMTP account
FALLBACK_FROM = os.getenv("EMAIL_USER", "")
FALLBACK_PASS = os.getenv("EMAIL_PASS", "")


def _ts() -> str:
    return datetime.now().strftime("%H:%M:%S")


# ── Colours ───────────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1F, 0x38)
C_ACCENT = RGBColor(0x6C, 0x63, 0xFF)
C_GREEN  = RGBColor(0x05, 0x96, 0x69)
C_RED    = RGBColor(0xDC, 0x26, 0x26)
C_AMBER  = RGBColor(0xD9, 0x77, 0x06)
C_GRAY   = RGBColor(0x4B, 0x55, 0x63)
C_LGRAY  = RGBColor(0x9C, 0xA3, 0xAF)


def _score_color(score: int) -> RGBColor:
    if score >= 75: return C_GREEN
    if score >= 50: return C_AMBER
    return C_RED


def _mini_bar(score: int, width: int = 12) -> str:
    filled = round(score / 100 * width)
    return "█" * filled + "░" * (width - filled)


def _yoe_from_resume(resume: dict) -> str:
    """Best-effort extract years of experience from resume JSON."""
    try:
        # Common keys the extraction agent might use
        for key in ("total_experience", "years_of_experience", "experience_years",
                    "total_work_experience", "work_experience_years"):
            val = resume.get(key)
            if val is not None:
                return str(val)
        # Count work experience entries as proxy
        exp = resume.get("work_experience") or resume.get("experience") or []
        if isinstance(exp, list) and exp:
            return f"~{len(exp)} role(s) listed"
    except Exception:
        pass
    return "N/A"


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _para(doc, text, bold=False, size=9, color=None, space_after=2, indent=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or C_DARK
    return p


def _section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    r.font.size  = Pt(8)
    r.font.bold  = True
    r.font.color.rgb = C_ACCENT
    # Underline via border-bottom on paragraph
    return p


def _bullet(doc, text, color=None, size=9):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color or C_DARK


# ── 1-page report builder ────────────────────────────────────────────────────

def build_report_docx(
    session: InterviewSession,
    evaluation: dict,
    chats: list[InterviewChat],
    candidate: Candidate,
) -> Path:
    doc = Document()

    # Narrow margins to fit everything on one page
    for sec in doc.sections:
        sec.top_margin    = Inches(0.45)
        sec.bottom_margin = Inches(0.45)
        sec.left_margin   = Inches(0.65)
        sec.right_margin  = Inches(0.65)

    # ── Extract data ──────────────────────────────────────────────────────────
    cand_m   = evaluation.get("candidate_metrics", {})
    bot_m    = evaluation.get("bot_metrics", {})
    sys_m    = evaluation.get("system_evaluation", {})
    flags    = evaluation.get("bias_flags", [])
    recs     = evaluation.get("recommendations", [])
    shrm     = evaluation.get("shrm_compliance_summary", {})
    overall  = evaluation.get("overall_analysis", "")

    cand_score = cand_m.get("overall_score", 0)
    bot_score  = bot_m.get("overall_score",  0)
    sys_score  = sys_m.get("overall_system_score", 0)

    screening  = candidate.screening_result or {}
    resume     = candidate.extracted_resume_json or {}
    role       = screening.get("jd_title") or screening.get("role") or "N/A"
    yoe        = _yoe_from_resume(resume)

    if cand_score >= 75:
        verdict, verdict_col, verdict_bg = "RECOMMENDED", C_GREEN, "D1FAE5"
    elif cand_score >= 55:
        verdict, verdict_col, verdict_bg = "CONSIDER WITH RESERVATIONS", C_AMBER, "FEF3C7"
    else:
        verdict, verdict_col, verdict_bg = "NOT RECOMMENDED", C_RED, "FEE2E2"

    # ══════════════════════════════════════════════════════════════════════════
    # HEADER  ─  title + verdict banner
    # ══════════════════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(1)
    tr = title_p.add_run("HR-Bot  ·  Candidate Evaluation Report")
    tr.font.size = Pt(15)
    tr.font.bold = True
    tr.font.color.rgb = C_ACCENT

    # Verdict line
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vp.paragraph_format.space_after = Pt(4)
    vr = vp.add_run(f"Hiring Verdict:  {verdict}")
    vr.font.size = Pt(10)
    vr.font.bold = True
    vr.font.color.rgb = verdict_col

    # ══════════════════════════════════════════════════════════════════════════
    # CANDIDATE INFO  +  SCORE SUMMARY  (2-col table)
    # ══════════════════════════════════════════════════════════════════════════
    _section_title(doc, "Candidate Profile")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(3.2)
    tbl.columns[1].width = Inches(3.2)

    left, right = tbl.rows[0].cells
    _set_cell_bg(left,  "F8F7FF")
    _set_cell_bg(right, "F0FDF4")

    # Left cell — candidate info
    for para in left.paragraphs:
        para.clear()
    def lcell(text, bold=False, color=None):
        p = left.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.bold = bold
        r.font.color.rgb = color or C_DARK

    lcell(f"Name:             {session.candidate_name or 'N/A'}", bold=True)
    lcell(f"Role Applied:     {role}")
    lcell(f"Experience:       {yoe}")
    lcell(f"Session ID:       #{session.id}")
    lcell(f"Date:             {datetime.now().strftime('%d %b %Y, %H:%M')}")
    lcell(f"Questions Asked:  {session.question_count} / {session.max_questions}")

    # Right cell — score summary
    for para in right.paragraphs:
        para.clear()
    def rcell(label, score):
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{label:<22}")
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = C_DARK
        r2 = p.add_run(f"{_mini_bar(score)}  {score}/100")
        r2.font.size = Pt(9)
        r2.font.name = "Courier New"
        r2.font.color.rgb = _score_color(score)

    rcell("Candidate Score", cand_score)
    rcell("AI Interviewer",  bot_score)
    rcell("System Score",    sys_score)
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(f"SHRM Verdict:  {shrm.get('overall_shrm_verdict', 'N/A')}")
    r.font.size = Pt(9)
    r.font.bold = True
    shrm_v = shrm.get("overall_shrm_verdict", "")
    r.font.color.rgb = C_GREEN if "Fully" in shrm_v else (C_AMBER if "Partially" in shrm_v else C_RED)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════════════════
    # STRENGTHS & WEAKNESSES  (side by side)
    # ══════════════════════════════════════════════════════════════════════════
    _section_title(doc, "Strengths & Areas for Improvement")

    cand_metric_map = [
        ("Communication",   cand_m.get("communication_clarity_score", 0)),
        ("Relevance",       cand_m.get("relevance_score", 0)),
        ("Technical",       cand_m.get("technical_competency_score", 0)),
        ("Confidence",      cand_m.get("confidence_conviction_score", 0)),
        ("Engagement",      cand_m.get("engagement_depth_score", 0)),
    ]
    sorted_metrics = sorted(cand_metric_map, key=lambda x: x[1], reverse=True)
    strengths  = [f"{l} ({s}/100)" for l, s in sorted_metrics if s >= 65]
    weaknesses = [f"{l} ({s}/100)" for l, s in sorted_metrics if s < 65]

    sw_tbl = doc.add_table(rows=1, cols=2)
    sw_tbl.style = "Table Grid"
    sw_l, sw_r = sw_tbl.rows[0].cells
    _set_cell_bg(sw_l, "F0FDF4")
    _set_cell_bg(sw_r, "FFF1F2")

    for para in sw_l.paragraphs: para.clear()
    for para in sw_r.paragraphs: para.clear()

    def sw_item(cell, text, color):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(f"• {text}")
        r.font.size = Pt(9)
        r.font.color.rgb = color

    p = sw_l.add_paragraph("✅ Strengths")
    p.runs[0].font.bold = True; p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = C_GREEN
    p.paragraph_format.space_after = Pt(1)
    for s in (strengths or ["No outstanding strengths identified"]):
        sw_item(sw_l, s, C_GREEN)

    p = sw_r.add_paragraph("⚠ Areas to Improve")
    p.runs[0].font.bold = True; p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = C_RED
    p.paragraph_format.space_after = Pt(1)
    for w in (weaknesses or ["No significant weaknesses identified"]):
        sw_item(sw_r, w, C_RED)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════════════════
    # EVALUATION METRICS SUMMARY  (compact bullets)
    # ══════════════════════════════════════════════════════════════════════════
    _section_title(doc, "Evaluation Metrics Summary")

    metrics_bullets = [
        (f"Candidate — Communication {cand_m.get('communication_clarity_score',0)}/100  |  "
         f"Technical {cand_m.get('technical_competency_score',0)}/100  |  "
         f"Confidence {cand_m.get('confidence_conviction_score',0)}/100  |  "
         f"Relevance {cand_m.get('relevance_score',0)}/100  |  "
         f"Engagement {cand_m.get('engagement_depth_score',0)}/100",
         C_DARK),
        (f"AI Interviewer — Question Quality {bot_m.get('question_quality_score',0)}/100  |  "
         f"Topic Coverage {bot_m.get('topic_coverage_score',0)}/100  |  "
         f"Adaptability {bot_m.get('adaptability_score',0)}/100  |  "
         f"Bias Risk {bot_m.get('bias_risk_score',0)}/100 (lower=better)",
         C_DARK),
        (f"System — Screening Accuracy {sys_m.get('screening_accuracy_score',0)}/100  |  "
         f"Interview Quality {sys_m.get('interview_quality_score',0)}/100  |  "
         f"Fairness {sys_m.get('fairness_transparency_score',0)}/100  |  "
         f"Candidate Experience {sys_m.get('candidate_experience_score',0)}/100",
         C_DARK),
        (f"Bias Flags: {len(flags)} detected" +
         (f"  [{', '.join(f['flag_type'] for f in flags[:3])}]" if flags else "  — None"),
         C_RED if flags else C_GREEN),
    ]
    for text, color in metrics_bullets:
        _bullet(doc, text, color=color, size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════════════════
    # OVERALL SUMMARY  (3 bullets)
    # ══════════════════════════════════════════════════════════════════════════
    _section_title(doc, "Overall Summary")

    # Split overall_analysis into up to 3 sentences as bullet points
    import re
    sentences = re.split(r'(?<=[.!?])\s+', overall.strip()) if overall else []
    # Fill to 3 bullets
    summary_points = sentences[:3] if len(sentences) >= 3 else sentences
    # Pad with generated bullet if fewer than 3
    cand_sum = cand_m.get("summary", "")
    bot_sum  = bot_m.get("summary", "")
    fallbacks = [cand_sum, bot_sum, f"Verdict: {verdict}. Candidate score {cand_score}/100."]
    i = 0
    while len(summary_points) < 3 and i < len(fallbacks):
        if fallbacks[i] and fallbacks[i] not in summary_points:
            summary_points.append(fallbacks[i])
        i += 1

    for pt in summary_points[:3]:
        _bullet(doc, pt.strip(), size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════════════════
    # RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════════════════
    if recs:
        _section_title(doc, "Recommendations")
        for rec in recs[:4]:          # cap at 4 to stay on one page
            _bullet(doc, rec, size=8)

    # ══════════════════════════════════════════════════════════════════════════
    # FOOTER line
    # ══════════════════════════════════════════════════════════════════════════
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(5)
    fr = footer_p.add_run(
        f"Generated by HR-Bot  ·  {datetime.now().strftime('%d %b %Y %H:%M')}  ·  Session #{session.id}"
    )
    fr.font.size = Pt(7)
    fr.font.color.rgb = C_LGRAY
    fr.italic = True

    # ── Save ─────────────────────────────────────────────────────────────────
    safe_name = (session.candidate_name or "candidate").replace(" ", "_").lower()
    fname = f"report_{safe_name}_session{session.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = REPORTS_DIR / fname
    doc.save(str(out_path))
    logger.info(f"[{_ts()}] 📄 Report saved → {out_path}")
    return out_path


# ── Email ─────────────────────────────────────────────────────────────────────

def _try_smtp_send(from_addr: str, password: str, msg) -> bool:
    """Attempt SMTP send. Returns True on success."""
    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(from_addr, password)
            server.send_message(msg)
        return True
    except Exception as exc:
        logger.warning(f"[{_ts()}] ⚠ SMTP send failed ({from_addr}): {exc}")
        return False


def send_report_email(report_path: Path, session: InterviewSession, evaluation: dict) -> bool:
    if not HR_EMAIL:
        logger.warning(f"[{_ts()}] ⚠ HR_EMAIL not set — skipping report email")
        return False

    cand_score = evaluation.get("candidate_metrics", {}).get("overall_score", 0)
    if cand_score >= 75:   verdict = "RECOMMENDED ✅"
    elif cand_score >= 55: verdict = "CONSIDER WITH RESERVATIONS ⚠️"
    else:                  verdict = "NOT RECOMMENDED ❌"

    subject = (
        f"[HR-Bot] Evaluation Report — {session.candidate_name} "
        f"| Session #{session.id} | {verdict}"
    )
    body = f"""Hello HR Team,

The AI-powered screening interview for the following candidate has been completed and evaluated.

Candidate:      {session.candidate_name}
Session ID:     #{session.id}
Date:           {datetime.now().strftime('%d %B %Y, %H:%M')}
Questions:      {session.question_count} / {session.max_questions}

EVALUATION SUMMARY
------------------
  Candidate Score:         {evaluation.get('candidate_metrics', {}).get('overall_score', 'N/A')} / 100
  AI Interviewer Score:    {evaluation.get('bot_metrics', {}).get('overall_score', 'N/A')} / 100
  System Score:            {evaluation.get('system_evaluation', {}).get('overall_system_score', 'N/A')} / 100

VERDICT: {verdict}

{evaluation.get('overall_analysis', '')}

Please find the full report attached as a Word document (.docx).

---
HR-Bot Evaluation System (auto-generated)
"""

    # Determine sender — prefer working fallback if primary likely has no app-pw
    sender_addr = FALLBACK_FROM or REPORT_FROM
    sender_pass = FALLBACK_PASS or REPORT_PASS

    if not sender_addr or not sender_pass:
        logger.error(f"[{_ts()}] ❌ No usable SMTP credentials — cannot send email")
        return False

    msg = MIMEMultipart()
    msg["From"]    = sender_addr
    msg["To"]      = HR_EMAIL
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))

    with open(report_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{report_path.name}"')
    msg.attach(part)

    # Try fallback first (known working app-password account)
    if FALLBACK_FROM and FALLBACK_PASS:
        if _try_smtp_send(FALLBACK_FROM, FALLBACK_PASS, msg):
            logger.info(f"[{_ts()}] 📧 Report emailed via fallback → {HR_EMAIL}")
            return True

    # Try primary (REPORT_FROM) if different
    if REPORT_FROM and REPORT_PASS and REPORT_FROM != FALLBACK_FROM:
        if _try_smtp_send(REPORT_FROM, REPORT_PASS, msg):
            logger.info(f"[{_ts()}] 📧 Report emailed via primary → {HR_EMAIL}")
            return True

    logger.error(f"[{_ts()}] ❌ All SMTP attempts failed — report not emailed")
    return False


# ── Agent class ───────────────────────────────────────────────────────────────

class FinalReportAgent:

    def generate(self, session_id: int, db: Session) -> dict:
        logger.info(f"[{_ts()}] ── FinalReportAgent START | session_id={session_id}")

        session = db.query(InterviewSession).filter(InterviewSession.id == session_id).first()
        if not session:
            return {"error": f"Session {session_id} not found."}

        evaluation = session.evaluation_result
        if not evaluation:
            return {
                "error": (
                    f"No evaluation data on session {session_id}. "
                    "Bias analysis must complete first (fires automatically after interview ends)."
                )
            }

        candidate = db.query(Candidate).filter(Candidate.id == session.candidate_id).first()
        if not candidate:
            return {"error": f"Candidate {session.candidate_id} not found."}

        chats = (
            db.query(InterviewChat)
            .filter(InterviewChat.session_id == session_id)
            .order_by(InterviewChat.question_number)
            .all()
        )

        try:
            report_path = build_report_docx(session, evaluation, chats, candidate)
        except Exception as exc:
            logger.error(f"[{_ts()}] ❌ Report build failed: {exc}", exc_info=True)
            return {"error": f"Report generation failed: {exc}"}

        session.report_path = str(report_path)
        db.commit()

        emailed = send_report_email(report_path, session, evaluation)

        logger.info(f"[{_ts()}] ✅ FinalReportAgent DONE | session={session_id} | emailed={emailed}")
        return {
            "status": "success",
            "session_id": session_id,
            "candidate_name": session.candidate_name,
            "report_path": str(report_path),
            "report_filename": report_path.name,
            "emailed_to": HR_EMAIL if emailed else None,
            "email_sent": emailed,
        }

    def generate_in_background(self, session_id: int) -> None:
        def _run():
            try:
                from app.database import SessionLocal
                db = SessionLocal()
                try:
                    result = self.generate(session_id=session_id, db=db)
                    if "error" in result:
                        logger.error(f"[{_ts()}] ❌ [BG] {result['error']}")
                    else:
                        logger.info(
                            f"[{_ts()}] ✅ [BG] {result.get('report_filename')} "
                            f"| emailed={result.get('email_sent')}"
                        )
                finally:
                    db.close()
            except Exception as exc:
                logger.error(f"[{_ts()}] ❌ [BG] Thread error: {exc}", exc_info=True)

        t = threading.Thread(target=_run, daemon=True, name=f"report-{session_id}")
        t.start()
        logger.info(f"[{_ts()}] 🚀 Report thread launched — session={session_id}")


final_report_agent = FinalReportAgent()
